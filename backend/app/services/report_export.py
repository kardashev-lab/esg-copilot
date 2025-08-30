import io
from typing import List, Dict, Any, Optional
from datetime import datetime
import json
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

class ReportExportService:
    """Service for exporting sustainability reports in various formats"""
    
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()
    
    def _setup_custom_styles(self):
        """Setup custom styles for reports"""
        
        # Custom title style
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1,  # Center alignment
            textColor=colors.darkblue
        )
        
        # Custom heading style
        self.heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20,
            textColor=colors.darkgreen
        )
        
        # Custom body style
        self.body_style = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            leading=14
        )
    
    async def export_report_pdf(self, report_data: Dict[str, Any], 
                              output_path: Optional[str] = None) -> bytes:
        """Export report to PDF format"""
        
        if output_path is None:
            output_path = f"sustainability_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Create PDF document
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        
        # Add title page
        story.extend(self._create_title_page(report_data))
        story.append(PageBreak())
        
        # Add table of contents
        story.extend(self._create_table_of_contents(report_data))
        story.append(PageBreak())
        
        # Add executive summary
        story.extend(self._create_executive_summary(report_data))
        story.append(PageBreak())
        
        # Add main content sections
        story.extend(self._create_environmental_section(report_data))
        story.append(PageBreak())
        
        story.extend(self._create_social_section(report_data))
        story.append(PageBreak())
        
        story.extend(self._create_governance_section(report_data))
        story.append(PageBreak())
        
        # Add appendices
        story.extend(self._create_appendices(report_data))
        
        # Build PDF
        doc.build(story)
        
        # Read the generated file and return bytes
        with open(output_path, 'rb') as f:
            pdf_bytes = f.read()
        
        # Clean up temporary file
        if os.path.exists(output_path):
            os.remove(output_path)
        
        return pdf_bytes
    
    async def export_report_docx(self, report_data: Dict[str, Any], 
                               output_path: Optional[str] = None) -> bytes:
        """Export report to DOCX format"""
        
        if output_path is None:
            output_path = f"sustainability_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Create Word document
        doc = Document()
        
        # Add title page
        self._add_title_page_docx(doc, report_data)
        
        # Add table of contents
        self._add_table_of_contents_docx(doc, report_data)
        
        # Add executive summary
        self._add_executive_summary_docx(doc, report_data)
        
        # Add main content sections
        self._add_environmental_section_docx(doc, report_data)
        self._add_social_section_docx(doc, report_data)
        self._add_governance_section_docx(doc, report_data)
        
        # Add appendices
        self._add_appendices_docx(doc, report_data)
        
        # Save document
        doc.save(output_path)
        
        # Read the generated file and return bytes
        with open(output_path, 'rb') as f:
            docx_bytes = f.read()
        
        # Clean up temporary file
        if os.path.exists(output_path):
            os.remove(output_path)
        
        return docx_bytes
    
    def _create_title_page(self, report_data: Dict[str, Any]) -> List:
        """Create title page for PDF"""
        
        elements = []
        
        # Company name
        company_name = report_data.get('company_name', 'Company Name')
        elements.append(Paragraph(company_name, self.title_style))
        elements.append(Spacer(1, 20))
        
        # Report title
        report_title = report_data.get('report_title', 'Sustainability Report')
        elements.append(Paragraph(report_title, self.title_style))
        elements.append(Spacer(1, 20))
        
        # Report period
        report_period = report_data.get('report_period', '2023')
        elements.append(Paragraph(f"Reporting Period: {report_period}", self.body_style))
        elements.append(Spacer(1, 20))
        
        # Publication date
        pub_date = datetime.now().strftime('%B %d, %Y')
        elements.append(Paragraph(f"Published: {pub_date}", self.body_style))
        elements.append(Spacer(1, 40))
        
        # Company information
        company_info = report_data.get('company_info', {})
        if company_info:
            info_text = f"""
            <b>Company Information:</b><br/>
            Industry: {company_info.get('industry', 'N/A')}<br/>
            Employees: {company_info.get('employees', 'N/A')}<br/>
            Revenue: {company_info.get('revenue', 'N/A')}<br/>
            Headquarters: {company_info.get('headquarters', 'N/A')}
            """
            elements.append(Paragraph(info_text, self.body_style))
        
        return elements
    
    def _create_table_of_contents(self, report_data: Dict[str, Any]) -> List:
        """Create table of contents for PDF"""
        
        elements = []
        
        elements.append(Paragraph("Table of Contents", self.heading_style))
        elements.append(Spacer(1, 20))
        
        # Define sections
        sections = [
            "Executive Summary",
            "Environmental Performance",
            "Social Performance", 
            "Governance Performance",
            "Appendices"
        ]
        
        for i, section in enumerate(sections, 1):
            elements.append(Paragraph(f"{i}. {section}", self.body_style))
            elements.append(Spacer(1, 6))
        
        return elements
    
    def _create_executive_summary(self, report_data: Dict[str, Any]) -> List:
        """Create executive summary section for PDF"""
        
        elements = []
        
        elements.append(Paragraph("Executive Summary", self.heading_style))
        elements.append(Spacer(1, 12))
        
        # Key highlights
        highlights = report_data.get('highlights', [])
        if highlights:
            elements.append(Paragraph("Key Highlights:", self.body_style))
            elements.append(Spacer(1, 6))
            
            for highlight in highlights:
                elements.append(Paragraph(f"• {highlight}", self.body_style))
                elements.append(Spacer(1, 3))
        
        elements.append(Spacer(1, 12))
        
        # Performance summary
        performance = report_data.get('performance_summary', {})
        if performance:
            elements.append(Paragraph("Performance Summary:", self.body_style))
            elements.append(Spacer(1, 6))
            
            for metric, value in performance.items():
                elements.append(Paragraph(f"• {metric}: {value}", self.body_style))
                elements.append(Spacer(1, 3))
        
        return elements
    
    def _create_environmental_section(self, report_data: Dict[str, Any]) -> List:
        """Create environmental performance section for PDF"""
        
        elements = []
        
        elements.append(Paragraph("Environmental Performance", self.heading_style))
        elements.append(Spacer(1, 12))
        
        # Climate action
        climate_data = report_data.get('environmental', {}).get('climate', {})
        if climate_data:
            elements.append(Paragraph("Climate Action", self.body_style))
            elements.append(Spacer(1, 6))
            
            # Create emissions table
            emissions_data = [
                ['Scope', 'Emissions (tCO2e)', 'Change from Previous Year']
            ]
            
            for scope, data in climate_data.get('emissions', {}).items():
                emissions_data.append([
                    scope,
                    str(data.get('value', 'N/A')),
                    str(data.get('change', 'N/A'))
                ])
            
            if len(emissions_data) > 1:
                emissions_table = Table(emissions_data)
                emissions_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(emissions_table)
                elements.append(Spacer(1, 12))
        
        # Energy efficiency
        energy_data = report_data.get('environmental', {}).get('energy', {})
        if energy_data:
            elements.append(Paragraph("Energy Efficiency", self.body_style))
            elements.append(Spacer(1, 6))
            
            for metric, value in energy_data.items():
                elements.append(Paragraph(f"• {metric}: {value}", self.body_style))
                elements.append(Spacer(1, 3))
        
        return elements
    
    def _create_social_section(self, report_data: Dict[str, Any]) -> List:
        """Create social performance section for PDF"""
        
        elements = []
        
        elements.append(Paragraph("Social Performance", self.heading_style))
        elements.append(Spacer(1, 12))
        
        # Diversity and inclusion
        diversity_data = report_data.get('social', {}).get('diversity', {})
        if diversity_data:
            elements.append(Paragraph("Diversity and Inclusion", self.body_style))
            elements.append(Spacer(1, 6))
            
            # Create diversity table
            diversity_table_data = [
                ['Metric', 'Value', 'Target']
            ]
            
            for metric, data in diversity_data.items():
                diversity_table_data.append([
                    metric,
                    str(data.get('value', 'N/A')),
                    str(data.get('target', 'N/A'))
                ])
            
            if len(diversity_table_data) > 1:
                diversity_table = Table(diversity_table_data)
                diversity_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(diversity_table)
                elements.append(Spacer(1, 12))
        
        # Labor rights
        labor_data = report_data.get('social', {}).get('labor', {})
        if labor_data:
            elements.append(Paragraph("Labor Rights", self.body_style))
            elements.append(Spacer(1, 6))
            
            for metric, value in labor_data.items():
                elements.append(Paragraph(f"• {metric}: {value}", self.body_style))
                elements.append(Spacer(1, 3))
        
        return elements
    
    def _create_governance_section(self, report_data: Dict[str, Any]) -> List:
        """Create governance performance section for PDF"""
        
        elements = []
        
        elements.append(Paragraph("Governance Performance", self.heading_style))
        elements.append(Spacer(1, 12))
        
        # Board diversity
        board_data = report_data.get('governance', {}).get('board', {})
        if board_data:
            elements.append(Paragraph("Board Diversity", self.body_style))
            elements.append(Spacer(1, 6))
            
            for metric, value in board_data.items():
                elements.append(Paragraph(f"• {metric}: {value}", self.body_style))
                elements.append(Spacer(1, 3))
        
        # Risk management
        risk_data = report_data.get('governance', {}).get('risk_management', {})
        if risk_data:
            elements.append(Paragraph("Risk Management", self.body_style))
            elements.append(Spacer(1, 6))
            
            for risk, details in risk_data.items():
                elements.append(Paragraph(f"• {risk}: {details}", self.body_style))
                elements.append(Spacer(1, 3))
        
        return elements
    
    def _create_appendices(self, report_data: Dict[str, Any]) -> List:
        """Create appendices section for PDF"""
        
        elements = []
        
        elements.append(Paragraph("Appendices", self.heading_style))
        elements.append(Spacer(1, 12))
        
        # Methodology
        methodology = report_data.get('methodology', '')
        if methodology:
            elements.append(Paragraph("Methodology", self.body_style))
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(methodology, self.body_style))
            elements.append(Spacer(1, 12))
        
        # Data sources
        data_sources = report_data.get('data_sources', [])
        if data_sources:
            elements.append(Paragraph("Data Sources", self.body_style))
            elements.append(Spacer(1, 6))
            
            for source in data_sources:
                elements.append(Paragraph(f"• {source}", self.body_style))
                elements.append(Spacer(1, 3))
        
        return elements
    
    def _add_title_page_docx(self, doc: Document, report_data: Dict[str, Any]):
        """Add title page to Word document"""
        
        # Title
        title = doc.add_heading(report_data.get('report_title', 'Sustainability Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company name
        company_name = report_data.get('company_name', 'Company Name')
        company_para = doc.add_paragraph(company_name)
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report period
        report_period = report_data.get('report_period', '2023')
        period_para = doc.add_paragraph(f"Reporting Period: {report_period}")
        period_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Publication date
        pub_date = datetime.now().strftime('%B %d, %Y')
        date_para = doc.add_paragraph(f"Published: {pub_date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
    
    def _add_table_of_contents_docx(self, doc: Document, report_data: Dict[str, Any]):
        """Add table of contents to Word document"""
        
        doc.add_heading("Table of Contents", level=1)
        
        sections = [
            "Executive Summary",
            "Environmental Performance",
            "Social Performance",
            "Governance Performance",
            "Appendices"
        ]
        
        for i, section in enumerate(sections, 1):
            doc.add_paragraph(f"{i}. {section}")
        
        doc.add_page_break()
    
    def _add_executive_summary_docx(self, doc: Document, report_data: Dict[str, Any]):
        """Add executive summary to Word document"""
        
        doc.add_heading("Executive Summary", level=1)
        
        # Key highlights
        highlights = report_data.get('highlights', [])
        if highlights:
            doc.add_paragraph("Key Highlights:")
            for highlight in highlights:
                doc.add_paragraph(highlight, style='List Bullet')
        
        # Performance summary
        performance = report_data.get('performance_summary', {})
        if performance:
            doc.add_paragraph("Performance Summary:")
            for metric, value in performance.items():
                doc.add_paragraph(f"• {metric}: {value}", style='List Bullet')
    
    def _add_environmental_section_docx(self, doc: Document, report_data: Dict[str, Any]):
        """Add environmental section to Word document"""
        
        doc.add_heading("Environmental Performance", level=1)
        
        # Climate action
        climate_data = report_data.get('environmental', {}).get('climate', {})
        if climate_data:
            doc.add_heading("Climate Action", level=2)
            
            # Create emissions table
            emissions_data = report_data.get('environmental', {}).get('climate', {}).get('emissions', {})
            if emissions_data:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                header_cells[0].text = 'Scope'
                header_cells[1].text = 'Emissions (tCO2e)'
                header_cells[2].text = 'Change from Previous Year'
                
                # Data rows
                for scope, data in emissions_data.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = scope
                    row_cells[1].text = str(data.get('value', 'N/A'))
                    row_cells[2].text = str(data.get('change', 'N/A'))
        
        # Energy efficiency
        energy_data = report_data.get('environmental', {}).get('energy', {})
        if energy_data:
            doc.add_heading("Energy Efficiency", level=2)
            for metric, value in energy_data.items():
                doc.add_paragraph(f"• {metric}: {value}", style='List Bullet')
    
    def _add_social_section_docx(self, doc: Document, report_data: Dict[str, Any]):
        """Add social section to Word document"""
        
        doc.add_heading("Social Performance", level=1)
        
        # Diversity and inclusion
        diversity_data = report_data.get('social', {}).get('diversity', {})
        if diversity_data:
            doc.add_heading("Diversity and Inclusion", level=2)
            
            # Create diversity table
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Metric'
            header_cells[1].text = 'Value'
            header_cells[2].text = 'Target'
            
            # Data rows
            for metric, data in diversity_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = metric
                row_cells[1].text = str(data.get('value', 'N/A'))
                row_cells[2].text = str(data.get('target', 'N/A'))
        
        # Labor rights
        labor_data = report_data.get('social', {}).get('labor', {})
        if labor_data:
            doc.add_heading("Labor Rights", level=2)
            for metric, value in labor_data.items():
                doc.add_paragraph(f"• {metric}: {value}", style='List Bullet')
    
    def _add_governance_section_docx(self, doc: Document, report_data: Dict[str, Any]):
        """Add governance section to Word document"""
        
        doc.add_heading("Governance Performance", level=1)
        
        # Board diversity
        board_data = report_data.get('governance', {}).get('board', {})
        if board_data:
            doc.add_heading("Board Diversity", level=2)
            for metric, value in board_data.items():
                doc.add_paragraph(f"• {metric}: {value}", style='List Bullet')
        
        # Risk management
        risk_data = report_data.get('governance', {}).get('risk_management', {})
        if risk_data:
            doc.add_heading("Risk Management", level=2)
            for risk, details in risk_data.items():
                doc.add_paragraph(f"• {risk}: {details}", style='List Bullet')
    
    def _add_appendices_docx(self, doc: Document, report_data: Dict[str, Any]):
        """Add appendices to Word document"""
        
        doc.add_heading("Appendices", level=1)
        
        # Methodology
        methodology = report_data.get('methodology', '')
        if methodology:
            doc.add_heading("Methodology", level=2)
            doc.add_paragraph(methodology)
        
        # Data sources
        data_sources = report_data.get('data_sources', [])
        if data_sources:
            doc.add_heading("Data Sources", level=2)
            for source in data_sources:
                doc.add_paragraph(f"• {source}", style='List Bullet')
    
    async def generate_sample_report_data(self) -> Dict[str, Any]:
        """Generate sample report data for testing"""
        
        return {
            'company_name': 'Sample Corporation',
            'report_title': 'Sustainability Report 2023',
            'report_period': '2023',
            'company_info': {
                'industry': 'Technology',
                'employees': '10,000+',
                'revenue': '$5 billion',
                'headquarters': 'San Francisco, CA'
            },
            'highlights': [
                'Reduced carbon emissions by 15% year-over-year',
                'Achieved 100% renewable energy for corporate operations',
                'Improved board diversity to 50% women',
                'Launched comprehensive supplier sustainability program'
            ],
            'performance_summary': {
                'Carbon Emissions': '45,000 tCO2e (-15%)',
                'Renewable Energy': '100%',
                'Board Diversity': '50% women',
                'Supplier Compliance': '95%'
            },
            'environmental': {
                'climate': {
                    'emissions': {
                        'Scope 1': {'value': '15,000', 'change': '-10%'},
                        'Scope 2': {'value': '30,000', 'change': '-20%'},
                        'Scope 3': {'value': '500,000', 'change': '-12%'}
                    }
                },
                'energy': {
                    'Renewable Energy': '100%',
                    'Energy Efficiency': '+25% improvement',
                    'Energy Intensity': '0.8 MWh/employee'
                }
            },
            'social': {
                'diversity': {
                    'Women in Workforce': {'value': '45%', 'target': '50%'},
                    'Women in Leadership': {'value': '40%', 'target': '50%'},
                    'Underrepresented Groups': {'value': '30%', 'target': '35%'}
                },
                'labor': {
                    'Employee Satisfaction': '85%',
                    'Training Hours': '45 hours/employee',
                    'Safety Incidents': '0.5 per 100 workers'
                }
            },
            'governance': {
                'board': {
                    'Independent Directors': '75%',
                    'Women Directors': '50%',
                    'Average Tenure': '6 years'
                },
                'risk_management': {
                    'ESG Risk Assessment': 'Completed annually',
                    'Climate Risk Integration': 'Fully integrated',
                    'Supply Chain Risk': 'Monitored quarterly'
                }
            },
            'methodology': 'This report follows GRI Standards 2021 and includes data from our integrated management systems.',
            'data_sources': [
                'Internal ESG data management system',
                'Third-party sustainability audits',
                'Employee surveys and feedback',
                'Supplier compliance reports'
            ]
        }
