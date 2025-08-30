import os
import uuid
from typing import List, Dict, Any, Optional
from datetime import datetime
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
import openpyxl
from pathlib import Path

from app.models.document import Document, DocumentType, DocumentStatus, DocumentCategory
from app.core.config import settings

class DocumentProcessor:
    def __init__(self):
        self.upload_dir = Path(settings.upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    async def process_document(self, file_content: bytes, filename: str, 
                             category: DocumentCategory, description: Optional[str] = None) -> Document:
        """Process uploaded document and extract text content"""
        
        # Generate unique ID and file path
        doc_id = str(uuid.uuid4())
        file_extension = Path(filename).suffix.lower()
        file_type = self._get_document_type(file_extension)
        
        # Save file
        file_path = self.upload_dir / f"{doc_id}{file_extension}"
        with open(file_path, "wb") as f:
            f.write(file_content)
        
        # Create document record
        document = Document(
            id=doc_id,
            filename=filename,
            file_type=file_type,
            category=category,
            description=description,
            file_path=str(file_path),
            file_size=len(file_content),
            status=DocumentStatus.PROCESSING,
            created_at=datetime.utcnow(),
            metadata={}
        )
        
        try:
            # Extract text content based on file type
            content = await self._extract_text(file_path, file_type)
            
            # Update document with extracted content
            document.metadata["extracted_content"] = content
            document.metadata["word_count"] = len(content.split())
            document.status = DocumentStatus.PROCESSED
            document.processed_at = datetime.utcnow()
            
        except Exception as e:
            document.status = DocumentStatus.FAILED
            document.metadata["error"] = str(e)
        
        return document
    
    def _get_document_type(self, file_extension: str) -> DocumentType:
        """Map file extension to document type"""
        extension_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".xlsx": DocumentType.XLSX,
            ".csv": DocumentType.CSV,
            ".txt": DocumentType.TXT
        }
        return extension_map.get(file_extension, DocumentType.TXT)
    
    async def _extract_text(self, file_path: Path, file_type: DocumentType) -> str:
        """Extract text content from different file types"""
        
        if file_type == DocumentType.PDF:
            return await self._extract_pdf_text(file_path)
        elif file_type == DocumentType.DOCX:
            return await self._extract_docx_text(file_path)
        elif file_type == DocumentType.XLSX:
            return await self._extract_xlsx_text(file_path)
        elif file_type == DocumentType.CSV:
            return await self._extract_csv_text(file_path)
        elif file_type == DocumentType.TXT:
            return await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        try:
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    async def _extract_xlsx_text(self, file_path: Path) -> str:
        """Extract text from XLSX file"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text += f"Sheet: {sheet_name}\n"
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Failed to extract XLSX text: {str(e)}")
    
    async def _extract_csv_text(self, file_path: Path) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        except Exception as e:
            raise Exception(f"Failed to extract CSV text: {str(e)}")
    
    async def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            raise Exception(f"Failed to extract TXT text: {str(e)}")
    
    def chunk_content(self, content: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split content into overlapping chunks for vector storage"""
        chunks = []
        start = 0
        
        while start < len(content):
            end = start + chunk_size
            chunk = content[start:end]
            
            # Try to break at sentence boundary
            if end < len(content):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > start + chunk_size // 2:
                    chunk = content[start:start + break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - overlap
            
            if start >= len(content):
                break
        
        return chunks
